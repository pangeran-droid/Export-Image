import os
from tkinter import Tk, Button, Label, filedialog, messagebox
from PIL import Image
from fpdf import FPDF
from docx import Document
from docx.shared import Inches

def pilih_folder():
    folder = filedialog.askdirectory()
    if folder:
        folder_label.config(text=folder)
        app.folder = folder

def export_pdf():
    try:
        files = ambil_gambar()
        pdf = FPDF(unit='mm', format='A4')
        for img in files:
            pdf.add_page()
            pdf.image(img, x=0, y=0, w=210, h=297)
        pdf.output("output.pdf", "F")
        messagebox.showinfo("Sukses", "PDF berhasil disimpan sebagai output.pdf")
    except Exception as e:
        messagebox.showerror("Gagal", str(e))

def export_word():
    try:
        files = ambil_gambar()
        doc = Document()
        for img in files:
            doc.add_picture(img, width=Inches(6))
            doc.add_paragraph("")
        doc.save("output.docx")
        messagebox.showinfo("Sukses", "Word berhasil disimpan sebagai output.docx")
    except Exception as e:
        messagebox.showerror("Gagal", str(e))

def export_png_panjang():
    try:
        files = ambil_gambar()
        images = [Image.open(img) for img in files]
        widths, heights = zip(*(i.size for i in images))
        max_width = max(widths)
        total_height = sum(heights)
        combined = Image.new('RGB', (max_width, total_height), (255, 255, 255))

        y_offset = 0
        for img in images:
            combined.paste(img, (0, y_offset))
            y_offset += img.height
        combined.save("combined_output.png")
        messagebox.showinfo("Sukses", "Gambar panjang disimpan sebagai combined_output.png")
    except Exception as e:
        messagebox.showerror("Gagal", str(e))

def export_webp():
    try:
        files = ambil_gambar()
        for img_path in files:
            img = Image.open(img_path)
            webp_path = os.path.splitext(img_path)[0] + '.webp'
            img.save(webp_path, 'WEBP', quality=80)
        messagebox.showinfo("Sukses", "Semua gambar telah dikonversi ke WebP")
    except Exception as e:
        messagebox.showerror("Gagal", str(e))

def ambil_gambar():
    folder = getattr(app, 'folder', None)
    if not folder:
        raise Exception("Silakan pilih folder gambar terlebih dahulu.")
    files = [os.path.join(folder, f) for f in os.listdir(folder)
             if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    if not files:
        raise Exception("Tidak ada gambar .jpg/.png ditemukan di folder.")
    return files

# GUI App
app = Tk()
app.title("Alat Konversi Gambar")
app.geometry("500x500")
app.config(bg="#2c3e50")

font_utama = ("Helvetica", 12)
font_judul = ("Helvetica", 16, "bold")
font_label = ("Helvetica", 10, "italic")

label_judul = Label(app, text="Alat Konversi Dokumen & Gambar", font=font_judul, bg="#2c3e50", fg="#ecf0f1")
label_judul.pack(pady=20)

label_folder = Label(app, text="Pilih folder berisi gambar:", font=font_utama, bg="#2c3e50", fg="#ecf0f1")
label_folder.pack(pady=(10, 5))

button_folder = Button(app, text="📁 Pilih Folder", command=pilih_folder, font=font_utama, bg="#3498db", fg="white", activebackground="#2980b9", relief="flat", padx=10, pady=5)
button_folder.pack()

folder_label = Label(app, text="Belum ada folder yang dipilih", font=font_label, bg="#2c3e50", fg="#ecf0f1")
folder_label.pack(pady=5)

label_ekspor = Label(app, text="Pilih format ekspor:", font=font_utama, bg="#2c3e50", fg="#ecf0f1")
label_ekspor.pack(pady=(20, 10))

button_pdf = Button(app, text="📝 Export ke PDF", width=30, command=export_pdf, font=font_utama, bg="#e74c3c", fg="white", activebackground="#c0392b", relief="flat", padx=10, pady=5)
button_pdf.pack(pady=5)

button_word = Button(app, text="📄 Export ke Word", width=30, command=export_word, font=font_utama, bg="#27ae60", fg="white", activebackground="#229954", relief="flat", padx=10, pady=5)
button_word.pack(pady=5)

button_png = Button(app, text="🖼️ Gabung PNG Panjang", width=30, command=export_png_panjang, font=font_utama, bg="#f39c12", fg="white", activebackground="#d35400", relief="flat", padx=10, pady=5)
button_png.pack(pady=5)

button_webp = Button(app, text="🌐 Konversi ke WebP", width=30, command=export_webp, font=font_utama, bg="#9b59b6", fg="white", activebackground="#8e44ad", relief="flat", padx=10, pady=5)
button_webp.pack(pady=5)

app.mainloop()